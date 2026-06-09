import os
import docx
from PyPDF2 import PdfReader
from PIL import Image
import io
import subprocess
import tempfile

class DocumentReader:
    def __init__(self):
        self.supported_formats = ['.docx', '.doc', '.pdf', '.txt', '.md', '.rtf', '.html']
    
    def read_docx(self, file_path):
        doc = docx.Document(file_path)
        content = []
        images = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                try:
                    image = Image.open(io.BytesIO(image_data))
                    images.append({
                        'content': image,
                        'filename': rel.target_ref.split('/')[-1] if '/' in rel.target_ref else 'image'
                    })
                except:
                    pass
        
        return {
            'text': '\n'.join(content),
            'images': images,
            'paragraphs': [p.text for p in doc.paragraphs]
        }
    
    def read_doc(self, file_path):
        success = False
        result = {'text': '', 'images': [], 'paragraphs': []}
        
        try:
            word = None
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            
            doc = word.Documents.Open(os.path.abspath(file_path), ReadOnly=True)
            
            try:
                text = doc.Content.Text
                paragraphs = [p.Range.Text.strip() for p in doc.Paragraphs if p.Range.Text.strip()]
                
                if text and len(text) > 0:
                    result['text'] = text
                    result['paragraphs'] = paragraphs
                    success = True
            finally:
                doc.Close(SaveChanges=0)
            
            if word:
                word.Quit()
                
        except Exception as e:
            if 'word' in dir() and word:
                try:
                    word.Quit()
                except:
                    pass
        
        if success:
            return result
        
        try:
            import textract
            text = textract.process(file_path).decode('utf-8', errors='ignore')
            if text and len(text.strip()) > 0:
                return {
                    'text': text,
                    'images': [],
                    'paragraphs': text.split('\n')
                }
        except:
            pass
        
        tools = [
            ('catdoc', [file_path]),
            ('antiword', [file_path]),
        ]
        
        for tool, args in tools:
            try:
                result_proc = subprocess.run(args, capture_output=True, text=True, encoding='utf-8', errors='ignore')
                if result_proc.returncode == 0 and result_proc.stdout and len(result_proc.stdout.strip()) > 0:
                    return {
                        'text': result_proc.stdout,
                        'images': [],
                        'paragraphs': result_proc.stdout.split('\n')
                    }
            except:
                pass
        
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name
            
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = 0
                
                doc = word.Documents.Open(os.path.abspath(file_path), ReadOnly=True)
                doc.SaveAs2(os.path.abspath(tmp_path), FileFormat=16)
                doc.Close(SaveChanges=0)
                word.Quit()
                
                return self.read_docx(tmp_path)
            except:
                pass
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        except:
            pass
        
        return self.read_unknown_format(file_path)
    
    def read_pdf(self, file_path):
        reader = PdfReader(file_path)
        content = []
        images = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                content.append(text)
            
            for image_file in page.images:
                try:
                    image = Image.open(io.BytesIO(image_file.data))
                    images.append({
                        'content': image,
                        'filename': image_file.name if image_file.name else 'image'
                    })
                except:
                    pass
        
        return {
            'text': '\n'.join(content),
            'images': images,
            'paragraphs': []
        }
    
    def read_txt(self, file_path):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return {
            'text': text,
            'images': [],
            'paragraphs': text.split('\n')
        }
    
    def read_md(self, file_path):
        return self.read_txt(file_path)
    
    def read_rtf(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            text = self.extract_text_from_rtf(content)
            return {
                'text': text,
                'images': [],
                'paragraphs': text.split('\n')
            }
        except:
            return self.read_txt(file_path)
    
    def read_html(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            text = self.extract_text_from_html(content)
            return {
                'text': text,
                'images': [],
                'paragraphs': text.split('\n')
            }
        except:
            return self.read_txt(file_path)
    
    def extract_text_from_rtf(self, rtf_content):
        import re
        text = re.sub(r'\\[a-z]+(\s+\d+)?', '', rtf_content)
        text = re.sub(r'\\[{}]', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    
    def extract_text_from_html(self, html_content):
        import re
        text = re.sub(r'<[^>]+>', '', html_content)
        text = re.sub(r'&[a-z]+;', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    
    def read_document(self, file_path):
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext not in self.supported_formats:
            return self.read_unknown_format(file_path)
        
        if ext == '.docx':
            return self.read_docx(file_path)
        elif ext == '.doc':
            return self.read_doc(file_path)
        elif ext == '.pdf':
            return self.read_pdf(file_path)
        elif ext == '.txt':
            return self.read_txt(file_path)
        elif ext == '.md':
            return self.read_md(file_path)
        elif ext == '.rtf':
            return self.read_rtf(file_path)
        elif ext == '.html':
            return self.read_html(file_path)
    
    def read_unknown_format(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            if text and len(text.strip()) > 0:
                return {
                    'text': text,
                    'images': [],
                    'paragraphs': text.split('\n')
                }
        except:
            pass
        
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)
            
            if header.startswith(b'\xD0\xCF\x11\xE0'):
                return self.read_doc(file_path)
        except:
            pass
        
        return {
            'text': '',
            'images': [],
            'paragraphs': []
        }
    
    def extract_sections(self, text, keywords_dict):
        sections = {}
        lines = text.split('\n')
        current_section = None
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            matched = False
            for section_name, keywords in keywords_dict.items():
                for keyword in keywords:
                    if keyword in line_stripped or line_stripped in keyword:
                        current_section = section_name
                        if current_section not in sections:
                            sections[current_section] = []
                        matched = True
                        break
                if matched:
                    break
            
            if current_section and not matched:
                sections[current_section].append(line_stripped)
        
        for section in sections:
            sections[section] = '\n'.join(sections[section])
        
        return sections
